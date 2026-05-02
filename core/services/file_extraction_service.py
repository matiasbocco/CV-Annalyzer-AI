"""
Multi-format CV text extraction.

Supported formats:
  .pdf   — pdfplumber (existing logic, moved here from pdf_service.py)
  .docx  — python-docx (paragraphs + table cells)
  .jpg / .jpeg / .png / .webp — OpenAI Vision API (gpt-4o-mini)

Public interface:
  extract_text_from_bytes(raw, filename) -> str   # primary; used by endpoints
  extract_text(file: UploadFile) -> str           # thin wrapper; reads then delegates
"""

import base64
import io
import re
from pathlib import Path

import pdfplumber
from fastapi import HTTPException, UploadFile
from openai import OpenAIError

from core.services.llm_service import client as _llm_client

# ── Constants ─────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".jpg", ".jpeg", ".png", ".webp"}

_IMAGE_MIME: dict[str, str] = {
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png":  "image/png",
    ".webp": "image/webp",
}

_VISION_SYSTEM_PROMPT = (
    "You are a CV parser. Extract all text content from this CV image exactly as it "
    "appears. Return only the extracted text, no commentary."
)

_MAX_WORDS      = 12_000
_TRUNCATION_SUFFIX = " [... CV truncated]"
_MAX_FILE_BYTES = 5 * 1024 * 1024   # 5 MB
_MAX_PDF_PAGES  = 3


# ── Magic-byte validation ─────────────────────────────────────────────────────

def _magic_ok(raw: bytes, ext: str) -> bool:
    """Return True if the file's leading bytes match the expected format."""
    if ext == ".pdf":
        return raw[:4] == b"%PDF"
    if ext == ".docx":
        return raw[:4] == b"PK\x03\x04"  # DOCX is a ZIP archive
    if ext in (".jpg", ".jpeg"):
        return raw[:3] == b"\xff\xd8\xff"
    if ext == ".png":
        return raw[:4] == b"\x89PNG"
    if ext == ".webp":
        return len(raw) >= 12 and raw[:4] == b"RIFF" and raw[8:12] == b"WEBP"
    return True


# ── Shared post-processing ────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Strip control chars, collapse whitespace, truncate to _MAX_WORDS."""
    # Remove null bytes and ASCII control chars; keep \t (0x09), \n (0x0a), \r (0x0d)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    words = text.split()
    if len(words) > _MAX_WORDS:
        return " ".join(words[:_MAX_WORDS]) + _TRUNCATION_SUFFIX
    return text


# ── Format-specific extractors ────────────────────────────────────────────────

def _extract_pdf(raw: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber; rejects CVs over 3 pages."""
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        page_count = len(pdf.pages)
        if page_count > _MAX_PDF_PAGES:
            raise HTTPException(
                400,
                f"El CV no puede tener más de {_MAX_PDF_PAGES} páginas. "
                f"Este archivo tiene {page_count} páginas.",
            )
        pages = [page.extract_text(layout=True) or "" for page in pdf.pages]
    return _clean("\n".join(pages))


def _extract_docx(raw: bytes) -> str:
    """Extract text from DOCX bytes using python-docx.

    Includes both paragraph text and table cell text because many CVs use
    tables for layout (skills matrix, education table, etc.).
    """
    import docx  # lazy import — only needed for DOCX files

    doc = docx.Document(io.BytesIO(raw))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return _clean("\n".join(parts))


async def _extract_image(raw: bytes, ext: str) -> str:
    """Extract CV text from an image using OpenAI Vision (gpt-4o-mini).

    Encodes the image as a base64 data URL and asks the model to return
    the raw text content of the CV. Uses 'high' detail for better OCR
    accuracy on dense CV layouts.
    """
    mime = _IMAGE_MIME[ext]
    data_url = f"data:{mime};base64,{base64.b64encode(raw).decode('utf-8')}"

    try:
        response = await _llm_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": _VISION_SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url, "detail": "high"},
                        }
                    ],
                },
            ],
            max_tokens=4096,
        )
    except OpenAIError as exc:
        raise HTTPException(502, "El servicio de extracción de imágenes no está disponible. Intentá más tarde.") from exc

    text = response.choices[0].message.content or ""
    return _clean(text)


# ── Public API ────────────────────────────────────────────────────────────────

async def extract_text_from_bytes(raw: bytes, filename: str) -> str:
    """Detect file type from extension and extract CV text.

    Raises HTTPException 400 for unsupported formats, oversized files, magic
    mismatch, corrupt files, or PDFs exceeding the page limit.
    """
    if len(raw) > _MAX_FILE_BYTES:
        raise HTTPException(400, "El archivo supera el límite de 5 MB.")

    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        accepted = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise HTTPException(
            400,
            f"Tipo de archivo no soportado: '{ext}'. Formatos aceptados: {accepted}",
        )

    if not _magic_ok(raw, ext):
        raise HTTPException(
            400,
            f"El contenido del archivo no coincide con su extensión '{ext}'. "
            "Verificá que el archivo no esté dañado.",
        )

    try:
        if ext == ".pdf":
            return _extract_pdf(raw)
        if ext == ".docx":
            return _extract_docx(raw)
        return await _extract_image(raw, ext)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(400, f"No se pudo procesar el archivo '{Path(filename).name}'.") from exc


async def extract_text(file: UploadFile) -> str:
    """Thin wrapper: read UploadFile bytes then delegate to extract_text_from_bytes."""
    raw = await file.read()
    return await extract_text_from_bytes(raw, file.filename or "")
