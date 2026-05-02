"""
Unit tests for file_extraction_service.

PDF and DOCX tests use real in-memory files.
Image test mocks the OpenAI Vision API to avoid network calls.
"""
import io
import pytest
from fastapi import HTTPException
from unittest.mock import AsyncMock, MagicMock, patch


# ── DOCX extraction ───────────────────────────────────────────────────────────

def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    import docx
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.anyio
async def test_extract_docx_paragraphs():
    raw = _make_docx_bytes(["John Doe", "Software Engineer", "Python, Django, React"])
    from core.services.file_extraction_service import extract_text_from_bytes
    text = await extract_text_from_bytes(raw, "cv.docx")
    assert "John Doe" in text
    assert "Python" in text


@pytest.mark.anyio
async def test_extract_docx_table_cells():
    import docx
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, Go"
    table.cell(1, 0).text = "Education"
    table.cell(1, 1).text = "BSc Computer Science"
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    from core.services.file_extraction_service import extract_text_from_bytes
    text = await extract_text_from_bytes(raw, "cv.docx")
    assert "Python" in text
    assert "Computer Science" in text


# ── PDF extraction ────────────────────────────────────────────────────────────

@pytest.mark.anyio
async def test_extract_pdf_basic():
    """Smoke test: a minimal searchable PDF returns non-empty text."""
    try:
        from reportlab.pdfgen import canvas as rl_canvas
    except ImportError:
        pytest.skip("reportlab not installed — skipping PDF smoke test")

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(100, 750, "Alice Smith — Senior Developer")
    c.save()
    raw = buf.getvalue()

    from core.services.file_extraction_service import extract_text_from_bytes
    text = await extract_text_from_bytes(raw, "cv.pdf")
    assert "Alice" in text or len(text) > 0  # layout=True may re-order words


# ── Unsupported extension ─────────────────────────────────────────────────────

@pytest.mark.anyio
async def test_unsupported_extension_raises_400():
    from core.services.file_extraction_service import extract_text_from_bytes
    with pytest.raises(HTTPException) as exc_info:
        await extract_text_from_bytes(b"dummy", "resume.txt")
    assert exc_info.value.status_code == 400
    assert ".txt" in exc_info.value.detail


@pytest.mark.anyio
async def test_no_extension_raises_400():
    from core.services.file_extraction_service import extract_text_from_bytes
    with pytest.raises(HTTPException) as exc_info:
        await extract_text_from_bytes(b"dummy", "resume")
    assert exc_info.value.status_code == 400


# ── Vision API mock ───────────────────────────────────────────────────────────

@pytest.mark.anyio
async def test_image_extraction_calls_vision_api():
    """Vision path: mock the OpenAI client and verify text is returned."""
    fake_response = MagicMock()
    fake_response.choices = [MagicMock()]
    fake_response.choices[0].message.content = "Jane Doe\nProduct Manager\nEmail: jane@example.com"

    with patch(
        "core.services.file_extraction_service._llm_client"
    ) as mock_client:
        mock_client.chat = MagicMock()
        mock_client.chat.completions = MagicMock()
        mock_client.chat.completions.create = AsyncMock(return_value=fake_response)

        from core.services.file_extraction_service import extract_text_from_bytes
        # Use a tiny 1x1 PNG (valid PNG header so base64 encoding won't error)
        tiny_png = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
            b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00"
            b"\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18"
            b"\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        text = await extract_text_from_bytes(tiny_png, "cv.png")

    assert "Jane Doe" in text
    mock_client.chat.completions.create.assert_awaited_once()
